"""DOCX extractor: headings, paragraphs and tables with structural provenance.

python-docx walks paragraphs and tables as separate collections, which loses the
document order.  Drilling programs interleave them ("Section 5.2 Hydraulics",
then a table of pump rates), so this extractor iterates the body XML and keeps
them in a single ordered stream - the paragraph index recorded in provenance is
therefore the *document order* index of the body element, which is what a reader
sees in Word's navigation pane.
"""

from __future__ import annotations

import re
from typing import Any

from ..__init__ import EXTRACTION_ENGINE_VERSION
from ..core.errors import ExtractionError
from .interfaces import DocumentComplexity, ExtractionContext, ProvenanceBuilder
from .normalized import (
    ExtractionMetadata,
    NormalizedDocument,
    Page,
    Paragraph,
    Section,
    Table,
    clean_text,
)

_HEADING_STYLE = re.compile(r"^(?:Heading|Title|Subtitle)(?:\s+(\d+))?$", re.IGNORECASE)
_NUMBERED_HEADING = re.compile(r"^\s*(\d+(?:\.\d+){0,3})\.?\s+(?P<title>.{2,160})$")


class DocxExtractor:
    """DOCX (Office Open XML) reader with structural provenance."""

    name = "docx"
    version = EXTRACTION_ENGINE_VERSION
    description = (
        "python-docx paragraph/table reader (headings, styles, tables, comments where available)."
    )

    def supports(self, context: ExtractionContext) -> tuple[bool, str]:
        if context.extension.lower() == ".docx":
            return True, "DOCX (python-docx)"
        if context.extension.lower() == ".doc":
            return (
                False,
                "legacy binary .doc is not supported by python-docx; convert or route via MinerU",
            )
        return False, f"extension {context.extension} is not DOCX"

    def probe(self, context: ExtractionContext) -> DocumentComplexity:
        complexity = DocumentComplexity()
        try:
            import docx

            document = docx.Document(str(context.path))
            paragraphs = len(document.paragraphs)
            tables = len(document.tables)
            complexity.pages = max(1, paragraphs // 40)
            complexity.table_count = tables
            complexity.has_text_layer = paragraphs > 0
            complexity.text_chars_per_page = sum(len(p.text) for p in document.paragraphs) / max(
                1, complexity.pages
            )
            if tables:
                complexity.reasons.append(f"{tables} table(s)")
        except Exception as exc:  # noqa: BLE001
            complexity.reasons.append(f"probe failed: {type(exc).__name__}: {exc}")
        return complexity

    def extract(
        self, context: ExtractionContext, provenance: ProvenanceBuilder
    ) -> NormalizedDocument:
        try:
            import docx
            from docx.oxml.ns import qn
        except Exception as exc:  # pragma: no cover
            raise ExtractionError(f"python-docx is not available: {exc}") from exc

        try:
            document = docx.Document(str(context.path))
        except Exception as exc:
            raise ExtractionError(
                f"Cannot open {context.filename}: {type(exc).__name__}: {exc}"
            ) from exc

        normalized = NormalizedDocument(
            metadata=ExtractionMetadata(
                filename=context.filename,
                path=str(context.path),
                sha256=context.sha256,
                extension=context.extension,
                mime_type=context.mime_type
                or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                size_bytes=context.size_bytes,
                engine=f"python-docx {_docx_version()}",
            )
        )
        core = document.core_properties
        normalized.metadata.extra["core_properties"] = {
            key: str(value)
            for key, value in {
                "title": getattr(core, "title", ""),
                "author": getattr(core, "author", ""),
                "created": getattr(core, "created", ""),
                "modified": getattr(core, "modified", ""),
                "revision": getattr(core, "revision", ""),
                "category": getattr(core, "category", ""),
                "subject": getattr(core, "subject", ""),
                "comments": getattr(core, "comments", ""),
            }.items()
            if value not in (None, "", "None")
        }
        if getattr(core, "created", None):
            normalized.metadata.document_date = str(getattr(core, "created", ""))

        body = document.element.body
        para_lookup = {id(p._p): p for p in document.paragraphs}
        table_lookup = {id(t._tbl): t for t in document.tables}

        index = 0
        page = 1  # DOCX has no intrinsic pagination; "page" = the sheet-like unit
        normalized.pages.append(Page(index=page, text="", label="Document body"))
        char_cursor = 0
        current_section: Section | None = None
        heading_stack: list[tuple[int, str]] = []
        table_counter = 0

        for child in body.iterchildren():
            tag = child.tag.split("}")[-1]
            if tag == "p":
                source = para_lookup.get(id(child))
                if source is None:
                    continue
                text = clean_text(source.text or "")
                style_name = ""
                try:
                    style_name = str(source.style.name or "") if source.style is not None else ""
                except Exception:  # noqa: BLE001 - unusual styles must not abort extraction
                    style_name = ""
                level = _heading_level(style_name, text)
                paragraph = Paragraph(
                    index=index,
                    text=text,
                    page=page,
                    block=index,
                    heading_level=level,
                    style=style_name,
                    char_start=char_cursor,
                    char_end=char_cursor + len(text),
                    provenance=provenance.docx(
                        heading=(current_section.label if current_section else None),
                        paragraph=index,
                        excerpt=text[:2000],
                        confidence=0.95,
                    ),
                )
                if level:
                    while heading_stack and heading_stack[-1][0] >= level:
                        heading_stack.pop()
                    heading_stack.append((level, text))
                    number_match = _NUMBERED_HEADING.match(text)
                    current_section = Section(
                        heading=text,
                        level=level,
                        page=page,
                        char_start=char_cursor,
                        number=(number_match.group(1) if number_match else ""),
                    )
                    normalized.sections.append(current_section)
                else:
                    paragraph.section = current_section.label if current_section else ""
                    if current_section is not None:
                        current_section.paragraph_indices.append(index)
                        current_section.char_end = char_cursor + len(text)
                if text:
                    normalized.paragraphs.append(paragraph)
                    char_cursor += len(text) + 2
                index += 1
            elif tag == "tbl":
                source_table = table_lookup.get(id(child))
                if source_table is None:
                    continue
                rows: list[list[str | None]] = []
                merged_cells = 0
                for row in source_table.rows:
                    cells: list[str | None] = []
                    seen: set[int] = set()
                    for cell in row.cells:
                        # python-docx repeats the same cell object for a merged span.  We keep
                        # the repeated value so the grid stays rectangular (a ragged table breaks
                        # every column-oriented check downstream) but count the duplicates, so the
                        # merge is visible in the record instead of silently re-drawn as data.
                        marker = id(cell._tc)
                        text = clean_text(cell.text or "")
                        if marker in seen:
                            merged_cells += 1
                        cells.append(text)
                        seen.add(marker)
                    rows.append(cells)
                rows = [row for row in rows if any((c or "").strip() for c in row)]
                if rows:
                    table = Table(
                        table_id=f"table{table_counter + 1}",
                        rows=rows,
                        caption=_table_caption(document, table_counter),
                        page=page,
                        anchor=f"body element {index}",
                        provenance=provenance.docx(
                            table=table_counter,
                            excerpt="\n".join(" | ".join(c or "" for c in r) for r in rows[:5])[
                                :2000
                            ],
                        ),
                        extra={"body_index": index, "merged_cells": merged_cells},
                    )
                    normalized.tables.append(table)
                    table_counter += 1
                index += 1

        normalized.pages[0].text = clean_text("\n\n".join(p.text for p in normalized.paragraphs))
        normalized.pages[0].char_end = char_cursor
        normalized.text = clean_text(
            "\n\n".join(
                [p.text for p in normalized.paragraphs] + [t.text() for t in normalized.tables]
            )
        )
        normalized.metadata.page_count = 1
        normalized.metadata.extra["body_elements"] = index
        normalized.metadata.extra["table_count"] = table_counter
        if not normalized.paragraphs and not normalized.tables:
            normalized.diagnostics.append("DOCX body contained no readable paragraphs or tables")
        if _has_tracked_changes(document, qn):
            normalized.diagnostics.append(
                "document contains tracked changes/revisions - read the revision state before relying on values"
            )
        return normalized


def _heading_level(style_name: str, text: str) -> int | None:
    match = _HEADING_STYLE.match(style_name or "")
    if match:
        return int(match.group(1) or 1)
    if style_name and style_name.lower() in {"title"}:
        return 1
    numbered = _NUMBERED_HEADING.match(text or "")
    if numbered and len(text) <= 120 and not text.rstrip().endswith("."):
        return numbered.group(1).count(".") + 1
    return None


def _table_caption(document: Any, table_index: int) -> str:
    try:
        for paragraph in document.paragraphs[:2000]:
            text = (paragraph.text or "").strip()
            if re.match(rf"(?i)^(table|tab\.?)\s*{table_index + 1}\b", text):
                return text[:160]
    except Exception:  # noqa: BLE001
        return ""
    return f"Table {table_index + 1}"


def _has_tracked_changes(document: Any, qn: Any) -> bool:
    try:
        for tag in ("w:ins", "w:del"):
            if document.element.body.findall(f".//{qn(tag)}"):
                return True
    except Exception:  # noqa: BLE001
        return False
    return False


def _docx_version() -> str:
    try:
        import docx

        return str(getattr(docx, "__version__", "?"))
    except Exception:  # noqa: BLE001
        return "?"


__all__ = ["DocxExtractor"]
