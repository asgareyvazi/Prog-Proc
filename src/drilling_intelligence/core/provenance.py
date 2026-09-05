"""Provenance: every extracted fact must point back to an exact source location
and be *verifiable* against the source file (master spec sections 17 and 95).

A provenance reference is a typed locator, not a free-text note:

    PDF     -> page / block / paragraph / section label
    XLSX    -> sheet / cell or range (+ whether the formula or the value was read)
    DOCX    -> heading path / paragraph index / table cell
    TEXT    -> line and character range

``verify_provenance`` re-opens the original file and re-reads the location, so
"where did this come from?" can be answered with a checkable answer rather than
a trust-me string.  This is what makes the audit trail real.
"""

from __future__ import annotations

import json
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

from .errors import ProvenanceError

_WHITESPACE = re.compile(r"\s+")


def normalise_text(text: str) -> str:
    """Whitespace/case-insensitive comparison form used for verification."""
    return _WHITESPACE.sub(" ", (text or "").replace("\u00a0", " ")).strip().lower()


# --------------------------------------------------------------------------- locators
@dataclass(frozen=True)
class SourceLocator:
    """Base class for typed source locators."""

    kind: str = ""

    def ref(self) -> str:
        raise NotImplementedError

    def to_dict(self) -> dict[str, Any]:
        payload = asdict(self)
        payload["locator_kind"] = self.kind
        return payload

    @staticmethod
    def from_dict(payload: dict[str, Any]) -> SourceLocator:
        kind = payload.get("locator_kind") or payload.get("kind")
        data = {k: v for k, v in payload.items() if k not in {"locator_kind", "kind"}}
        cls = {
            "pdf": PdfLocator,
            "excel": ExcelLocator,
            "docx": DocxLocator,
            "text": TextLocator,
            "unknown": UnknownLocator,
        }.get(str(kind), UnknownLocator)
        valid = {f for f in cls.__dataclass_fields__ if f != "kind"}  # type: ignore[attr-defined]
        return cls(**{k: v for k, v in data.items() if k in valid})  # type: ignore[arg-type]


@dataclass(frozen=True)
class PdfLocator(SourceLocator):
    """A location in a PDF.  ``page`` is 1-based (what a human reads)."""

    page: int | None = None
    block: int | None = None
    paragraph: int | None = None
    section: str | None = None
    table: int | None = None
    bbox: tuple[float, float, float, float] | None = None
    kind: str = "pdf"

    def ref(self) -> str:
        bits: list[str] = []
        if self.page is not None:
            bits.append(f"Page {self.page}")
        if self.section:
            bits.append(f"Section {self.section}")
        if self.block is not None:
            bits.append(f"Block {self.block}")
        if self.paragraph is not None:
            bits.append(f"Paragraph {self.paragraph}")
        if self.table is not None:
            bits.append(f"Table {self.table}")
        return " > ".join(bits) if bits else "PDF (location not recorded)"


@dataclass(frozen=True)
class ExcelLocator(SourceLocator):
    """A cell or a rectangular range on a named sheet."""

    sheet: str = ""
    cell: str | None = None
    range_: str | None = None
    read: str = "value"  # value | formula
    row: int | None = None
    column: int | None = None
    kind: str = "excel"

    def ref(self) -> str:
        bits = []
        if self.sheet:
            bits.append(f"Sheet: {self.sheet}")
        if self.cell:
            bits.append(f"Cell: {self.cell}")
        elif self.range_:
            bits.append(f"Range: {self.range_}")
        if self.read == "formula":
            bits.append("formula")
        return " > ".join(bits) if bits else "XLSX (location not recorded)"


@dataclass(frozen=True)
class DocxLocator(SourceLocator):
    heading: str | None = None
    heading_path: tuple[str, ...] | None = None
    paragraph: int | None = None
    table: int | None = None
    row: int | None = None
    column: int | None = None
    kind: str = "docx"

    def ref(self) -> str:
        bits: list[str] = []
        if self.heading:
            bits.append(f"Heading: {self.heading}")
        if self.paragraph is not None:
            bits.append(f"Paragraph {self.paragraph}")
        if self.table is not None:
            part = f"Table {self.table}"
            if self.row is not None:
                part += f" > Row {self.row}"
                if self.column is not None:
                    part += f" > Cell {self.column}"
            bits.append(part)
        return " > ".join(bits) if bits else "DOCX (location not recorded)"


@dataclass(frozen=True)
class TextLocator(SourceLocator):
    """Line/character range in a plain-text or markdown file."""

    line_start: int | None = None
    line_end: int | None = None
    char_start: int | None = None
    char_end: int | None = None
    section: str | None = None
    kind: str = "text"

    def ref(self) -> str:
        bits: list[str] = []
        if self.section:
            bits.append(f"Section {self.section}")
        if self.line_start is not None:
            span = f"Lines {self.line_start}"
            if self.line_end is not None and self.line_end != self.line_start:
                span += f"-{self.line_end}"
            bits.append(span)
        elif self.char_start is not None:
            span = f"Chars {self.char_start}"
            if self.char_end is not None:
                span += f"-{self.char_end}"
            bits.append(span)
        return " > ".join(bits) if bits else "Text (location not recorded)"


@dataclass(frozen=True)
class UnknownLocator(SourceLocator):
    note: str = ""
    kind: str = "unknown"

    def ref(self) -> str:
        return f"Location unknown ({self.note})" if self.note else "Location unknown"


# --------------------------------------------------------------------------- record
@dataclass(frozen=True)
class Provenance:
    """A traceable reference to where a piece of information came from."""

    document_id: str
    filename: str
    locator: SourceLocator
    document_version_id: str | None = None
    parser: str = ""
    excerpt: str = ""
    #: sha256 of the source file the excerpt was read from (tamper/eager check).
    source_sha256: str = ""
    #: MinerU/OCR confidence or heuristic confidence for the field itself.
    confidence: float | None = None

    @property
    def ref(self) -> str:
        """Canonical human reference, e.g. ``DDR.xlsx > Sheet: Daily Report > Range: B14:F18``."""
        return f"{self.filename} > {self.locator.ref()}"

    def to_dict(self) -> dict[str, Any]:
        return {
            "document_id": self.document_id,
            "document_version_id": self.document_version_id,
            "filename": self.filename,
            "parser": self.parser,
            "excerpt": self.excerpt,
            "source_sha256": self.source_sha256,
            "confidence": self.confidence,
            "locator": self.locator.to_dict(),
        }

    def to_json(self) -> str:
        return json.dumps(self.to_dict(), sort_keys=True, ensure_ascii=False)

    @classmethod
    def from_dict(cls, payload: dict[str, Any]) -> Provenance:
        locator = payload.get("locator")
        return cls(
            document_id=str(payload.get("document_id", "")),
            filename=str(payload.get("filename", "")),
            locator=SourceLocator.from_dict(dict(locator or {})),
            document_version_id=payload.get("document_version_id"),
            parser=str(payload.get("parser", "")),
            excerpt=str(payload.get("excerpt", "")),
            source_sha256=str(payload.get("source_sha256", "")),
            confidence=payload.get("confidence"),
        )

    def __str__(self) -> str:  # pragma: no cover - cosmetic
        return self.ref


def provenance_list(payloads: list[dict[str, Any]] | None) -> list[Provenance]:
    if not payloads:
        return []
    return [Provenance.from_dict(p) for p in payloads]


def as_provenance(value: Provenance | dict[str, Any] | None) -> Provenance | None:
    if value is None:
        return None
    if isinstance(value, Provenance):
        return value
    return Provenance.from_dict(value)


# --------------------------------------------------------------------------- verification
@dataclass
class VerificationResult:
    """Outcome of re-reading a provenance reference against the source file."""

    provenance: Provenance
    status: str = "UNVERIFIED"  # MATCH | MISMATCH | UNREADABLE | NOT_CHECKABLE
    detail: str = ""
    current_excerpt: str = ""

    @property
    def ok(self) -> bool:
        return self.status == "MATCH"

    def to_dict(self) -> dict[str, Any]:
        return {
            "status": self.status,
            "detail": self.detail,
            "current_excerpt": self.current_excerpt,
            "ref": self.provenance.ref,
            "source_sha256_expected": self.provenance.source_sha256,
        }


def _excel_read(source: Path, locator: ExcelLocator) -> str:
    from openpyxl import load_workbook  # imported lazily: optional dependency at use time

    wb = load_workbook(source, data_only=locator.read != "formula", read_only=False)
    try:
        if locator.sheet not in wb.sheetnames:
            raise ProvenanceError(f"Sheet {locator.sheet!r} no longer exists", file=str(source))
        ws = wb[locator.sheet]
        if locator.cell:
            return str(ws[locator.cell].value or "")
        if locator.range_:
            rows = []
            for row in ws[locator.range_]:
                rows.append("\t".join("" if c.value is None else str(c.value) for c in row))
            return "\n".join(rows)
        raise ProvenanceError("Excel provenance records neither a cell nor a range", file=str(source))
    finally:
        wb.close()


def _pdf_read(source: Path, locator: PdfLocator) -> str:
    import pymupdf  # PyMuPDF

    with pymupdf.open(source) as doc:
        if locator.page is None or locator.page > doc.page_count or locator.page < 1:
            raise ProvenanceError(f"Page {locator.page} out of range", file=str(source))
        page = doc[locator.page - 1]
        if locator.table is not None:
            tables = page.find_tables().tables
            if locator.table >= len(tables):
                raise ProvenanceError(f"Table {locator.table} not found on page {locator.page}", file=str(source))
            rows = tables[locator.table].extract()
            return "\n".join("\t".join("" if c is None else str(c) for c in row) for row in rows)
        blocks = page.get_text("blocks")
        if locator.block is not None and 0 <= locator.block < len(blocks):
            return str(blocks[locator.block][4])
        return page.get_text()


def _docx_read(source: Path, locator: DocxLocator) -> str:
    import docx  # python-docx

    document = docx.Document(str(source))
    if locator.table is not None:
        if locator.table >= len(document.tables):
            raise ProvenanceError(f"Table {locator.table} missing", file=str(source))
        table = document.tables[locator.table]
        if locator.row is not None and locator.row < len(table.rows):
            cells = table.rows[locator.row].cells
            if locator.column is not None and locator.column < len(cells):
                return cells[locator.column].text
            return " | ".join(c.text for c in cells)
        return " | ".join(" | ".join(c.text for c in r.cells) for r in table.rows)
    if locator.paragraph is not None:
        if locator.paragraph >= len(document.paragraphs):
            raise ProvenanceError(f"Paragraph {locator.paragraph} missing", file=str(source))
        return document.paragraphs[locator.paragraph].text
    return "\n".join(p.text for p in document.paragraphs)


def _text_read(source: Path, locator: TextLocator) -> str:
    text = source.read_text(encoding="utf-8", errors="replace")
    if locator.char_start is not None:
        end = locator.char_end if locator.char_end is not None else len(text)
        return text[locator.char_start:end]
    if locator.line_start is not None:
        lines = text.splitlines()
        start = max(1, locator.line_start)
        end = min(len(lines), locator.line_end or start)
        return "\n".join(lines[start - 1 : end])
    return text


def verify_provenance(
    source: Path | str | None,
    provenance: Provenance,
    *,
    require_hash: bool = True,
    fuzzy: float = 0.9,
) -> VerificationResult:
    """Re-read the recorded location and compare with the stored excerpt.

    ``require_hash`` makes a changed file an immediate ``MISMATCH`` (the
    document on disk is not the document that produced the fact).  When the
    hash matches but text differs slightly (normalisation, extraction version)
    a similarity threshold decides ``MATCH`` vs ``MISMATCH``.
    """

    import difflib

    if source is None:
        return VerificationResult(provenance, "NOT_CHECKABLE", "No source path supplied")
    path = Path(source)
    if not path.exists():
        return VerificationResult(provenance, "UNREADABLE", f"Source file missing: {path}")

    if require_hash and provenance.source_sha256:
        from .hashing import sha256_file

        actual = sha256_file(path)
        if actual != provenance.source_sha256:
            return VerificationResult(
                provenance,
                "MISMATCH",
                "Source file hash changed since extraction - the recorded excerpt may no longer be valid",
            )

    locator = provenance.locator
    try:
        if isinstance(locator, ExcelLocator):
            current = _excel_read(path, locator)
        elif isinstance(locator, PdfLocator):
            current = _pdf_read(path, locator)
        elif isinstance(locator, DocxLocator):
            current = _docx_read(path, locator)
        elif isinstance(locator, TextLocator):
            current = _text_read(path, locator)
        else:
            return VerificationResult(provenance, "NOT_CHECKABLE", "Locator type cannot be re-read")
    except ProvenanceError as exc:
        return VerificationResult(provenance, "UNREADABLE", str(exc))
    except Exception as exc:  # noqa: BLE001 - pragma: no cover; defensive, third-party readers
        return VerificationResult(provenance, "UNREADABLE", f"{type(exc).__name__}: {exc}")

    expected_norm = normalise_text(provenance.excerpt)
    current_norm = normalise_text(current)
    if not expected_norm:
        return VerificationResult(provenance, "NOT_CHECKABLE", "No excerpt stored to compare", current_excerpt=current[:200])
    if expected_norm == current_norm or expected_norm in current_norm:
        return VerificationResult(provenance, "MATCH", "", current_excerpt=current[:200])
    ratio = difflib.SequenceMatcher(None, expected_norm, current_norm).ratio()
    if ratio >= fuzzy:
        return VerificationResult(provenance, "MATCH", f"Similar (ratio {ratio:.3f})", current_excerpt=current[:200])
    return VerificationResult(
        provenance,
        "MISMATCH",
        f"Excerpt differs (ratio {ratio:.3f}); re-extract the document",
        current_excerpt=current[:200],
    )


def provenances_from_rows(rows: list[dict[str, Any]]) -> list[Provenance]:
    out: list[Provenance] = []
    for row in rows:
        try:
            out.append(Provenance.from_dict(row))
        except Exception:  # noqa: BLE001 - tolerate legacy/corrupt rows rather than fail a whole listing
            continue
    return out


__all__ = [
    "DocxLocator",
    "ExcelLocator",
    "PdfLocator",
    "Provenance",
    "SourceLocator",
    "TextLocator",
    "UnknownLocator",
    "VerificationResult",
    "as_provenance",
    "normalise_text",
    "provenance_list",
    "provenances_from_rows",
    "verify_provenance",
]
