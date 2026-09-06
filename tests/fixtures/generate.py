"""Realistic test corpus builders.

Tests must run against genuine files, not mocks: an extractor that "works"
against a fake byte string tells you nothing about PDF table detection or
merged spreadsheet cells.  Every value written here is mirrored in
``GROUND_TRUTH`` so assertions are exact rather than fuzzy.

The files are generated on demand into ``tmp_path`` (see ``tests/conftest.py``)
and are also usable by ``tools/make_sample_corpus.py`` for a manual demo.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

#: Values the fixtures assert and that tests can check end to end.  ``files`` is
#: the set of corpus files that state the value; ``unit`` is what the *document*
#: writes (unit conversion is verified separately in the unit tests).
GROUND_TRUTH: dict[str, dict[str, Any]] = {
    "mud_weight": {
        "value": 10.2,
        "unit": "ppg",
        "files": (
            "mud_report_well-a3.xlsx",
            "well_a3_program_rev12.pdf",
            "daily_drilling_report_well-a3.docx",
        ),
    },
    "depth_md": {
        "value": 10125.0,
        "unit": "ft",
        "files": ("mud_report_well-a3.xlsx", "daily_drilling_report_well-a3.docx"),
    },
    "depth_tvd": {"value": 9850.0, "unit": "ft", "files": ("mud_report_well-a3.xlsx",)},
    "casing_shoe_depth": {
        "value": 8500.0,
        "unit": "ft",
        "files": ("daily_drilling_report_well-a3.docx",),
    },
    "fracture_gradient": {"value": 15.8, "unit": "ppg", "files": ("well_a3_program_rev12.pdf",)},
    "equivalent_mud_weight": {"value": 15.8, "unit": "ppg", "files": ("mud_report_well-a3.xlsx",)},
    "pore_pressure_gradient": {
        "value": 0.465,
        "unit": "psi/ft",
        "files": ("mud_report_well-a3.xlsx", "well_a3_program_rev12.pdf"),
    },
    "npt_hours": {"value": 18.5, "unit": "h", "files": ("daily_drilling_report_well-a3.docx",)},
    "hole_size_in": {
        "value": 12.25,
        "unit": "in",
        "files": ("well_a3_program_rev12.pdf", "daily_drilling_report_well-a3.docx"),
    },
    "casing_size_in": {"value": 9.625, "unit": "in", "files": ("well_a3_program_rev12.pdf",)},
    "rop": {"value": 42.0, "unit": "ft/hr", "files": ("daily_drilling_report_well-a3.docx",)},
    "flow_rate": {"value": 900.0, "unit": "gpm", "files": ("well_a3_program_rev12.pdf",)},
    "mud_volume_bbl": {
        "value": 1450.0,
        "unit": "bbl",
        "files": ("mud_report_well-a3.xlsx", "daily_drilling_report_well-a3.docx"),
    },
    "wob": {
        "value": 32000.0,
        "unit": "lbf",
        "files": ("well_a3_program_rev12.pdf", "daily_drilling_report_well-a3.docx"),
    },
    "torque": {"value": 18400.0, "unit": "ft-lbf", "files": ("well_a3_program_rev12.pdf",)},
    "rpm": {
        "value": 120.0,
        "unit": "rpm",
        "files": ("well_a3_program_rev12.pdf", "daily_drilling_report_well-a3.docx"),
    },
    "date_iso": {"value": "2025-06-14", "unit": "", "files": ("mud_report_well-a3.xlsx",)},
}

#: Values that must NOT be reported for a field name, with the reason.  A wrong
#: value in the knowledge base is worse than a missing one (master spec section 46),
#: so these negative cases are asserted by the tests exactly like the positive ones.
NEGATIVE_TRUTH: list[dict[str, object]] = [
    {
        "field": "mud_weight",
        "forbidden": [11.4],
        "file": "well_a3_program_rev12.pdf",
        "why": "the program says 'Do not exceed 11.4 ppg': that is a limit, not the design mud weight",
    },
    {
        "field": "mud_weight",
        "forbidden": [15.8],
        "file": "mud_report_well-a3.xlsx",
        "why": "15.8 ppg is the EMW row (dynamic density) and belongs to equivalent_mud_weight",
    },
]


def build_mud_report_xlsx(path: Path) -> Path:
    """A mud report workbook: summary sheet with label/value pairs, a data sheet
    with formulas, merged headers and one hidden sheet (all real-world noise)."""
    from openpyxl import Workbook
    from openpyxl.styles import Font

    workbook = Workbook()
    summary = workbook.active
    summary.title = "Summary"
    summary["A1"] = "ACME DRILLING - WELL A-3 MUD REPORT"
    summary["A1"].font = Font(bold=True, size=14)
    summary.merge_cells("A1:D1")
    rows = [
        ("Well", "A-3", "", ""),
        ("Field", "North Cormorant", "", ""),
        ("Report date", "2025-06-14", "", ""),
        ("Revision", "3", "", ""),
        ("MD (ft)", 10125.0, "ft", "Measured depth at report time"),
        ("TVD (ft)", 9850.0, "ft", "True vertical depth"),
        ("Mud weight (ppg)", 10.2, "ppg", "Rheometer reading at 500/300 rpm"),
        ("Plastic viscosity", 18.0, "cP", ""),
        ("Yield point", 12.0, "lb/100ft2", ""),
        ("Gel strength 10s", 6.0, "lb/100ft2", ""),
        ("Chloride (mg/l)", 18500.0, "mg/l", "Seawater baseline 19000"),
        ("EMW (ppg)", 15.8, "ppg", "Equivalent mud weight during circulation"),
        ("Pore pressure gradient", 0.465, "psi/ft", ""),
        ("Total mud volume (bbl)", 1450.0, "bbl", "Active system volume"),
    ]
    for index, row in enumerate(rows, start=3):
        for column, value in enumerate(row, start=1):
            summary.cell(row=index, column=column, value=value)
    summary.cell(row=9, column=2).number_format = '0.0" ppg"'

    data = workbook.create_sheet("Daily Tests")
    data.append(["Slip", "Time", "MW in (ppg)", "MW out (ppg)", "Visc (cP)", "Sand (pct)", "Notes"])
    readings = [
        ("1st", "06:00", 10.15, 10.2, 18.0, 0.1, "Normal drilling"),
        ("2nd", "12:00", 10.2, 10.22, 19.0, 0.2, "Pumped pill"),
        ("3rd", "18:00", 10.22, 10.2, 18.5, 0.1, "Stable"),
    ]
    for reading in readings:
        data.append(reading)
    data.append(["", "Average MW out", "=AVERAGE(C2:C5)", "", "", "", ""])
    data["C6"].value = "=AVERAGE(D2:D5)"

    hidden = workbook.create_sheet("Calibration")
    hidden.sheet_state = "hidden"
    hidden.append(["Device", "Last calibration", "Result"])
    hidden.append(["Mud balance", "2025-05-30", "OK"])
    path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(path)
    return path


def build_ddr_docx(path: Path) -> Path:
    """A daily drilling report as Word: headings, narrative, a time breakdown table."""
    from docx import Document
    from docx.shared import Pt

    document = Document()
    document.core_properties.title = "Daily Drilling Report - Well A-3"
    document.core_properties.author = "R. Halvorsen (Drilling Supervisor)"
    document.add_heading("Daily Drilling Report - Well A-3", level=1)
    document.add_paragraph("Date: 14 June 2025   Report status: APPROVED")
    document.add_paragraph("Well: A-3   Section: 12 1/4 in intermediate   Trip number: 14")
    document.add_heading("Drilling summary", level=2)
    document.add_paragraph(
        "Drilled from 9,780 ft MD to 10,125 ft MD on 12 1/4 in bit 13 (IADC 1-1-1). "
        "Average rate of penetration was 42 ft/hr over the 8.25 h of drilling time, "
        "with mud weight maintained at 10.2 ppg and 1,450 bbl total system volume. "
        "WOB averaged 32,000 lbf with 120 rpm and 520 psi standpipe pressure."
    )
    document.add_paragraph(
        "Maximum allowable annular surface pressure was limited to 1,850 psi by the 9 5/8 in casing shoe at 8,500 ft MD."
    )
    document.add_heading("NPT", level=2)
    document.add_paragraph(
        "Stuck bit for 6.5 h while back reaming at 9,940 ft; 12.0 h lost to tripping equipment failure (top drive hose). Total NPT 18.5 h."
    )
    document.add_heading("Time breakdown", level=2)
    table = document.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    for cell, text in zip(header, ("Activity", "Hours", "Code"), strict=True):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    for activity, hours, code in (
        ("Drilling", 8.25, "DRILL"),
        ("Tripping", 14.0, "TRIP"),
        ("Circulating", 1.5, "CIRC"),
        ("Maintenance", 0.25, "MAINT"),
        ("NPT - stuck bit", 6.5, "NPT"),
        ("NPT - equipment", 12.0, "NPT"),
    ):
        row = table.add_row().cells
        row[0].text = activity
        row[1].text = f"{hours:.2f}"
        row[2].text = code
    document.add_paragraph(
        "Prepared by R. Halvorsen, Drilling Supervisor. Reviewed by the company man 15 June 2025."
    )
    style = document.styles["Normal"]
    style.font.size = Pt(10)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path


def build_program_pdf(path: Path) -> Path:
    """A drilling program revision as a text-layer PDF with a ruled hydraulics table.

    Written with PyMuPDF so that the page genuinely contains selectable text and
    vector table borders - the conditions the PDF extractor is designed for.
    """
    try:
        import pymupdf as fitz
    except ImportError:  # pragma: no cover - PyMuPDF < 1.24
        import fitz

    document = fitz.open()
    page = document.new_page(width=595, height=842)
    left = 50.0

    def line(text: str, y: float, size: float = 10.0, bold: bool = False) -> None:
        font = "hebo" if bold else "helv"
        page.insert_text((left, y), text, fontsize=size, fontname=font, color=(0, 0, 0))

    page.draw_rect(fitz.Rect(40, 30, 555, 62), color=(0.4, 0.4, 0.4), width=1)
    line("NORTH CORMORANT FIELD - DRILLING PROGRAM", 50, 14, True)
    line(
        "Well A-3   |   12 1/4 in intermediate section   |   Revision 12   |   Status: APPROVED",
        70,
        9,
    )
    y = 95.0
    for text in (
        "1. Objectives and constraints",
        "Drill the 12 1/4 in section from 8,500 ft MD (9 5/8 in casing shoe) to 10,450 ft MD.",
        "The casing shoe test on the previous well gave a fracture gradient of 15.8 ppg.",
        "Design mud weight is 10.2 ppg with an ECD target of 10.6 ppg.",
        "Pore pressure gradient is 0.465 psi/ft to total depth; TVD at section TD is 10,180 ft.",
        "Hole size is 12 1/4 in with a 5 in drill string. Maximum allowable annular surface",
        "pressure is calculated from the 15.8 ppg fracture gradient at the shoe.",
        "",
        "2. Hydraulics (see table 1)",
        "Pump rate 900 gpm at 520 psi standpipe pressure. Bit nozzle pressure loss 178 psi.",
        "Annular velocity in the 12 1/4 in hole is 142 ft/min, which is above the 120 ft/min",
        "minimum cut-off transport requirement. Torque averaged 18,400 ft-lbf in the previous",
        "section with a WOB of 32,000 lbf at 120 rpm.",
        "",
        "3. Contingency",
        "If losses exceed 10 bbl/hr at the shoe, reduce ECD by tripping to a lower pump rate",
        "and notify the well engineer before proceeding. Do not exceed 11.4 ppg mud weight.",
    ):
        if text:
            line(text, y, 10, text.endswith("(see table 1)") or text.startswith(("1.", "2.", "3.")))
            if text.startswith(("1.", "2.", "3.")):
                page.insert_text((left, y), "", fontsize=10)
        y += 16.0

    # Ruled table: drawn cell borders make it detectable by table finding.
    top = y + 12
    headers = ("Depth (ft MD)", "MW (ppg)", "Hydro (psi)", "FG (ppg)", "MAASP (psi)")
    data = [
        ("8,500", "10.2", "4,508", "15.8", "1,850"),
        ("9,500", "10.2", "5,038", "15.6", "1,420"),
        ("10,450", "10.4", "5,640", "15.6", "1,050"),
    ]
    column_widths = (110.0, 85.0, 90.0, 80.0, 90.0)
    row_height = 20.0
    x = left
    line("Table 1 - Annular pressure and MAASP summary", top - 8, 9, True)
    for row_index, row in enumerate((headers, *data)):
        y_cell = top + 6 + row_index * row_height
        x = left
        for column_index, cell in enumerate(row):
            width = column_widths[column_index]
            page.draw_rect(
                fitz.Rect(x, y_cell, x + width, y_cell + row_height),
                color=(0.2, 0.2, 0.2),
                width=0.6,
            )
            page.insert_text(
                (x + 5, y_cell + 13),
                cell,
                fontsize=9,
                fontname=("hebo" if row_index == 0 else "helv"),
            )
            x += width
    page.insert_text(
        (left, top + 6 + 4 * row_height + 20), "End of program extract - page 1 of 1", fontsize=8
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    document.close()
    return path


def build_npt_csv(path: Path) -> Path:
    """Time/NPT data as exported from a reporting system (CSV is a first-class input)."""
    payload = """Event No,Well,Date,Activity,NPT Hours,Code,Description
1,A-3,2025-06-13,Drilling,6.5,NPT-STUCK,Back reaming to free stuck bit at 9940 ft MD
2,A-3,2025-06-14,Tripping,12.0,NPT-EQUIP,Top drive hose failure during trip out
3,A-3,2025-06-12,Circulating,0.0,NPT-OTHER,No NPT recorded
4,B-11,2025-04-02,Drilling,22.25,NPT-STUCK,Washout below motor; tripping for replacement
"""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(payload, encoding="utf-8")
    return path


def build_lesson_txt(path: Path) -> Path:
    """A lessons-learned note: plain text, where line provenance is everything."""
    payload = """LESSON LEARNED - NORTH CORMORANT FIELD - LL-2025-014
Subject: Stuck bit after tripping for washout
Status: ACCEPTED
Date: 2025-06-20
Author: Drilling Engineer (K. Adeyemi)

1. Situation
   Well A-3 was sidetracked from 9,100 ft MD. A washout below the motor was
   identified on the second trip. While back reaming from 9,940 ft MD the bit
   became stuck for 6.5 hours; 12.0 hours of additional NPT followed a top
   drive hose failure.

2. What worked
   Reaming with 10.2 ppg mud and 900 gpm freed the string without a jarring
   campaign. Differential sticking was ruled out by the pressure test.

3. What did not work
   The hole cleaning survey was skipped on the previous well (B-11), where the
   same 12 1/4 in section cost 22.25 hours of NPT for the identical reason.
   Bedding was the root cause, and the offset data was available.

4. Recommendation
   Make the cuttings bed survey mandatory for sections below 9,000 ft MD where
   the fracture gradient at the shoe limits ECD to less than 15.8 ppg.
   Design mud weight window: 10.2 ppg to 11.4 ppg.
"""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(payload, encoding="utf-8")
    return path


def build_scanned_pdf(path: Path) -> Path:
    """A PDF with *no* text layer: the honest trigger for the MinerU fallback."""
    try:
        import pymupdf as fitz
    except ImportError:  # pragma: no cover - PyMuPDF < 1.24
        import fitz

    document = fitz.open()
    page = document.new_page(width=595, height=842)
    image = build_placeholder_scan()
    page.insert_image(fitz.Rect(20, 20, 575, 822), stream=image)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    document.close()
    return path


def build_placeholder_scan() -> bytes:
    """A tiny PNG with dark marks, enough to make a page look scanned."""
    import struct
    import zlib

    width, height = 64, 88
    rows = bytearray()
    for y in range(height):
        rows.append(0)  # filter: none
        for x in range(width):
            ink = (
                0
                if (10 <= y <= 14 and 6 <= x <= 58)
                or (20 + (y // 6) * 6 <= y <= 22 + (y // 6) * 6 and 6 <= x <= 50)
                else 255
            )
            rows.extend((ink, ink, ink))

    def chunk(tag: bytes, payload: bytes) -> bytes:
        return (
            struct.pack(">I", len(payload))
            + tag
            + payload
            + struct.pack(">I", zlib.crc32(tag + payload) & 0xFFFFFFFF)
        )

    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(bytes(rows), 9))
        + chunk(b"IEND", b"")
    )


BUILDERS = {
    "mud_report_well-a3.xlsx": build_mud_report_xlsx,
    "daily_drilling_report_well-a3.docx": build_ddr_docx,
    "well_a3_program_rev12.pdf": build_program_pdf,
    "npt_summary_2025-06.csv": build_npt_csv,
    "lesson_learned_ll-2025-014.txt": build_lesson_txt,
    "scanned_well_b11_report.pdf": build_scanned_pdf,
}


def build_corpus(root: Path | str, *, include_scan: bool = True) -> dict[str, Path]:
    """Write the fixture corpus into ``root`` and return ``{filename: path}``."""
    root = Path(root)
    root.mkdir(parents=True, exist_ok=True)
    written: dict[str, Path] = {}
    for name, builder in BUILDERS.items():
        if name.startswith("scanned_") and not include_scan:
            continue
        written[name] = builder(root / name)
    return written


__all__ = [
    "BUILDERS",
    "GROUND_TRUTH",
    "build_corpus",
    "build_ddr_docx",
    "build_lesson_txt",
    "build_mud_report_xlsx",
    "build_npt_csv",
    "build_placeholder_scan",
    "build_program_pdf",
    "build_scanned_pdf",
]
