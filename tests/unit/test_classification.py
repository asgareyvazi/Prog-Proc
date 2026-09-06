"""Classification: a label must be earned by evidence and must admit its own doubts.

These are the behaviours the UI shows the engineer: the label, a confidence that means
something, the matched text behind it, and an explicit weak/ambiguous note instead of a
confident-looking wrong answer.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import pytest

from drilling_intelligence.classification import (
    TAXONOMY,
    DeterministicClassifier,
    DocumentClassification,
)
from drilling_intelligence.classification.rules import MIN_MARGIN, NOISE_SCORE, WEAK_CONFIDENCE


@dataclass
class _Document:
    text: str
    extension: str


def _read(path: Path) -> _Document:
    """Read a corpus file the way the extractors would, without the parser overhead."""
    suffix = path.suffix.lstrip(".")
    if suffix == "pdf":
        import pymupdf

        with pymupdf.open(path) as doc:
            return _Document("\n".join(page.get_text() for page in doc), suffix)
    if suffix == "docx":
        import docx

        return _Document("\n".join(p.text for p in docx.Document(str(path)).paragraphs), suffix)
    if suffix in {"xlsx", "xlsm"}:
        from openpyxl import load_workbook

        book = load_workbook(path, data_only=True)
        rows = (
            " ".join(str(v) for v in row if v is not None)
            for sheet in book
            for row in sheet.iter_rows(values_only=True)
        )
        return _Document("\n".join(rows), suffix)
    return _Document(path.read_text(encoding="utf-8"), suffix)


@pytest.fixture(scope="module")
def classifier() -> DeterministicClassifier:
    return DeterministicClassifier()


def classify(classifier: DeterministicClassifier, filename: str, text: str, extension: str):
    return classifier.classify(filename=filename, text=text, extension=extension)


def test_the_real_corpus_lands_on_the_right_type(classifier, corpus_dir) -> None:
    expected = {
        "daily_drilling_report_well-a3.docx": DocumentClassification.DDR,
        "well_a3_program_rev12.pdf": DocumentClassification.DRILLING_PROGRAM,
        "mud_report_well-a3.xlsx": DocumentClassification.MUD_REPORT,
        "npt_summary_2025-06.csv": DocumentClassification.NPT,
        "lesson_learned_ll-2025-014.txt": DocumentClassification.LESSON_LEARNED,
    }
    for filename, classification in expected.items():
        document = _read(corpus_dir / filename)
        result = classify(classifier, filename, document.text, document.extension)
        assert result.classification is classification, (filename, result.scores, result.notes)
        assert not result.weak and not result.ambiguous, result.to_dict()
        assert result.confidence > 0.5, result.to_dict()


def test_separators_do_not_change_the_verdict(classifier, corpus_dir) -> None:
    document = _read(corpus_dir / "daily_drilling_report_well-a3.docx")
    verdicts = {
        classify(classifier, name, document.text, "docx").classification
        for name in (
            "daily_drilling_report_well-a3.docx",
            "daily-drilling-report-well-a3.docx",
            "Daily Drilling Report.docx",
        )
    }
    assert verdicts == {DocumentClassification.DDR}


def test_every_decision_carries_the_text_that_produced_it(classifier, corpus_dir) -> None:
    document = _read(corpus_dir / "well_a3_program_rev12.pdf")
    result = classify(classifier, "well_a3_program_rev12.pdf", document.text, "pdf")
    assert result.evidence, "a classification without evidence is not auditable"
    for item in result.evidence:
        assert item.source in {"filename", "content", "extension", "structure"}
        if item.source == "content":
            assert item.matched_text.lower() in document.text.lower()
    payload = result.to_dict()
    assert payload["evidence"] and payload["confidence"] == round(result.confidence, 4)


def test_scores_cover_the_taxonomy_and_the_winner_leads(classifier, corpus_dir) -> None:
    document = _read(corpus_dir / "mud_report_well-a3.xlsx")
    result = classify(classifier, "mud_report_well-a3.xlsx", document.text, "xlsx")
    assert set(result.scores) >= {signature.classification.value for signature in TAXONOMY}
    assert result.classification is DocumentClassification.MUD_REPORT
    assert result.scores[result.classification.value] == pytest.approx(max(result.scores.values()))
    runner_up = result.runner_up
    assert runner_up is not None
    assert (
        result.scores[result.classification.value] - runner_up[1] >= MIN_MARGIN
    ) or result.ambiguous


def test_a_scan_without_text_says_it_cannot_read_it(classifier) -> None:
    # Exactly the scanned-PDF case.  A vague name plus no readable text is not evidence
    # for a type, so the honest answer is OTHER - with the reason recorded so the UI can
    # offer "run OCR/MinerU" instead of showing a confident-looking wrong label.
    result = classify(classifier, "scanned_well_b11_report.pdf", "", "pdf")
    assert result.classification is DocumentClassification.OTHER
    assert result.confidence == 0.0 and result.weak
    assert any("OCR/MinerU required" in note for note in result.notes)


def test_a_descriptive_name_may_hint_when_the_text_is_unreadable(classifier) -> None:
    result = classify(classifier, "daily_drilling_report_well-a3.docx", "", "docx")
    assert result.classification is DocumentClassification.DDR
    # Filename-only evidence is capped: a name can suggest, never prove.
    assert result.confidence <= 0.45
    assert any("OCR/MinerU required" in note for note in result.notes)


def test_thin_evidence_is_other_not_the_closest_guess(classifier) -> None:
    result = classify(classifier, "report.pdf", "notes about the weather", "pdf")
    assert result.classification is DocumentClassification.OTHER and result.weak
    assert result.confidence <= NOISE_SCORE
    assert (
        classify(classifier, "IMG_2040.jpg", "", "jpg").classification
        is DocumentClassification.OTHER
    )
    assert WEAK_CONFIDENCE > 0


def test_authority_tier_follows_type_and_status(classifier, corpus_dir) -> None:
    document = _read(corpus_dir / "well_a3_program_rev12.pdf")
    approved = classifier.classify(
        filename="well_a3_program_rev12.pdf",
        text=document.text,
        extension="pdf",
        declared_status="APPROVED",
    )
    assert approved.authority_tier == "approved_drilling_program"
    superseded = classifier.classify(
        filename="well_a3_program_rev11.pdf",
        text=document.text,
        extension="pdf",
        declared_status="APPROVED",
        is_current=False,
    )
    assert superseded.authority_tier == "previous_revision"


def test_only_taxonomy_members_are_acceptable() -> None:
    # The contract any future LLM classifier has to satisfy: free text is not a label.
    assert DeterministicClassifier.validate("mud_report") is DocumentClassification.MUD_REPORT
    assert (
        DeterministicClassifier.validate("Drilling Program")
        is DocumentClassification.DRILLING_PROGRAM
    )
    assert DeterministicClassifier.validate("probably a report?") is None
    assert DeterministicClassifier.validate(None) is None
